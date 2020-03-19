import gi
gi.require_version('Gtk', '3.0')
from gi.repository import Gtk
from datetime import date
from datetime import datetime
import os.path
import platform, subprocess
import shutil
import re
from docx import Document
today = date.today()
def document_open():
    document = Document("./datos/plantilla.docx")
    print(document.paragraphs[3].text)
    document.save('./datos/new_document.docx')

def formated_namber(value):
    formated = "{:,}".format(value) + " Gs"
    return formated 

def today_file(read_format):
    today = date.today()
    filepath = "./datos/" + str(today) + ".txt"
    from_file = open(filepath,read_format)

class PhotocopyManager:            
    current_log_file = None
    total = 0
    current_user = "Marta"
    current_print_type = "black_and_white"
    initial_value_in_the_box = 0
    previous_added_price = 0
    def init_logger(self):
        print("Init")
        if not os.path.isdir("./datos"):
            os.mkdir("./datos")

        today = date.today()
        label_date = builder.get_object("label_date")
        label_date.set_text(str(today))
        filepath = "./datos/" + str(today) + ".txt" 
        if os.path.isfile(filepath):
            readed_file = open("./datos/"+str(today)+".txt","r")
            lines = readed_file.readlines()
    current_print_type = "black_and_white"
    initial_value_in_the_box = 0
    previous_added_price = 0
    def init_logger(self):
        print("Init")
        if not os.path.isdir("./datos"):
            os.mkdir("./datos")

        today = date.today()
        label_date = builder.get_object("label_date")
        label_date.set_text(str(today))
        filepath = "./datos/" + str(today) + ".txt" 
        if os.path.isfile(filepath):
            readed_file = open("./datos/"+str(today)+".txt","r")
            lines = readed_file.readlines()
            position = lines[0].find("TOTAL: ")
            #print(lines[0][position+7])
            offset1=position+7
            offset2=len(lines[0])-4
            result=lines[0][offset1:offset2]
            intresult=result.replace(",", "")
            new_string = ''.join(e for e in intresult if e.isalnum())
            re.sub('[^A-Za-z0-9]+', '', new_string)
            self.total=int(new_string)
        else:
            new_file = open("./datos/"+str(today)+".txt","w+")
            new_file.write("Fecha: ")
            new_file.write(str(today) + "                                ")
            new_file.write("TOTAL: 0 Gs\n")
            new_file.write("Hora                 Cantidad                Tipo\n")
            new_file.close()


class Handler:
    manager = None 
    total = 0
    dialog = None
    def __init__(self, manager):
        print("init") 
        document_open()
    def init(self):
        #manager.init_logger()    
        self.manager = manager 
        self.total = manager.total
        self.update_total_label()
        self.update_half_total_label()
        button1 = builder.get_object("rb_1")
        button1.connect("toggled",self.on_radio_button_marta_select)

        button2 = builder.get_object("rb_2")
        button3 = builder.get_object("rb_3")
        button2.connect("toggled",self.on_radio_button_oski_select)
        button3.connect("toggled",self.on_radio_button_pancha_select)

        print_black_and_white_radio = builder.get_object("black_white_radio_button")
        print_color_radio = builder.get_object("color_print_radio_button")
        print_black_and_white_radio.connect("toggled",self.radio_button_black_white_pressed)
        print_color_radio.connect("toggled",self.radio_button_color_pressed)
    def radio_button_black_white_pressed(self,button):
        self.manager.current_print_type = "black_and_white"
    def radio_button_color_pressed(self, button):
        self.manager.current_print_type = "color"

    def update_total_label(self):
        label_total = builder.get_object("label_total")
        label_total.set_text(formated_namber(self.total))
        label_total_in_the_box = builder.get_object("label_total_in_the_box")
        label_total_in_the_box.set_text(formated_namber(self.manager.initial_value_in_the_box + self.total)) 
        self.update_half_total_label()
    def print_total_to_inform_file(self):
        from_file = today_file("r") 
        line = from_file.readline()
        # make any changes to line here
        line = "Fecha: "
        line += str(today) + "                               "
        line += "TOTAL: " + formated_namber(self.total) + "\n"
        to_file = today_file("w") 
        to_file.write(line)
        shutil.copyfileobj(from_file, to_file)
        
    
    def update_half_total_label(self):
        label_half_total = builder.get_object("label_halft_total")  
        label_half_total.set_text(formated_namber(self.total/2)) 

    def print_total(self, price, data_type):
        current_log_file = open("./datos/"+str(today)+".txt","a")
        self.total = self.total + price
        self.update_total_label() 
        current_time = datetime.now().strftime("%H:%M:%S")
        formated_price =  "{:,}".format(price)
        current_log_file.write(current_time + 
                "                " + formated_price + 
                "                " + data_type + "\n" )
        current_log_file.close() 
        self.print_total_to_inform_file()
        self.manager.previous_added_price = price

    ###########################################
    ############       Buttons      ###########
    ###########################################
    def button_show_extract_clicked(self , button):
        print("Extracts")

    def button_print_service_clicked(self, button):
        print("Print service")
        black = True
        if(self.manager.current_print_type == "black_and_white"):
            self.print_total(1000,"Impresion en Blanco y Negro")
        elif (self.manager.current_print_type == "color"):
            self.print_total(2500,"Impresion a Color")

    def button_input_value_in_box_pressed(self , button):
        input_box = builder.get_object("input_value_in_box")
        value = input_box.get_text()
        value = int(value)
        self.manager.initial_value_in_the_box = value
        self.update_total_label()

    def button_SET_pressed(self, button):
        self.print_total(8000,"Certificado Contribuyente / No Contributente")

    def button_add_ID_count(self, button):
        self.print_total(1000,"Fotocopia de Cedula")
    def button_add_photocopie_count(self, button):
        self.print_total(500,"Fotocopia Simple Blanco y Negro")
    def button_curriculum_pressed(self, button):
        self.print_total(10000,"Curriculum")
    def button_judment_pressed(self, button):
        self.print_total(9000,"Antecedente Judicial")
    def button_folder_pressed(self, button):
        self.print_total(2000,"Carpeta")
    def button_plastic_pressed(self, button):
        self.print_total(500,"Folio")

    def button_undo_pressed(self, button):
        print("undo") 
        opend_file = today_file("r")
        lines = opend_file.readlines()
        opend_file.close()
        new_file = today_file("w")
        for line in range(0,(len(lines)-1)):
            new_file.write(lines[line])
        new_file.close()
        self.total -= self.manager.previous_added_price
        self.print_total_to_inform_file()
        self.update_total_label()
    def button_retire_pressed(self, button):
        print("retire") 
        self.dialog = builder.get_object("retire_dialog")
        response = self.dialog.run()
    
    def on_dialog_delete_event(self, dialog, event):
        dialog.hide()
        return True

    def button_retire_accept_pressed(self, button):
        input_retire_mount = builder.get_object("dialog_mount_input")
        input_value = input_retire_mount.get_text()
        print(input_value)
        print("accept")
        
        today = date.today()
        out_log_file= open("./datos/"+str(today)+"_out"+".txt","w+")
        out_log_file.write(self.manager.current_user+": ") 
        out_log_file.write(input_value) 
        out_log_file.write("\n") 
        out_log_file.close()
        self.dialog.hide()

    def on_radio_button_select(self, widget , data=None):
        print("radio button changed")
    def on_radio_button_marta_select(self,widget):
        print("marta")
        self.manager.current_user = "Marta"

    def on_radio_button_pancha_select(self,widget):
        print("pacha")
        self.manager.current_user = "Pancha"

    def on_radio_button_oski_select(self,widget):
        print("oski")
        self.manager.current_user = "Oski"

    def button_print_pressed(self, button):
        print("printing")
        today = date.today()
        if platform.system() == 'Windows':    # Windows
            filepath = "datos/" + str(today) + ".txt" 
            relative_path = os.path.abspath(filepath) 
            os.startfile(relative_path, "print") 

    def button_cancel_clicked(self, button):
        print("cancel")
        self.dialog.hide()

    def button_show_data_pressed(self, button):
        today = date.today()
        if platform.system() == 'Windows':    # Windows
            filepath = "datos/" + str(today) + ".txt" 
            relative_path = os.path.abspath(filepath) 
            os.startfile(relative_path)
        else:
            filepath = "./datos/" + str(today) + ".txt" 
            subprocess.call(('xdg-open', filepath))

    def insert_text_bold(self, document, paragraph_index  ,title , value):
        document.paragraphs[paragraph_index].text = ''
        run = document.paragraphs[paragraph_index].add_run(title)
        run.font.bold = True
        run = document.paragraphs[paragraph_index].add_run(value)
        run.font.bold = True
    

    def modify_table(self, document):
        print("table")
        document.tables[0].cell(1,0).text = "new value"

    def on_button_generate_pressed(self, button):
        print("generate")
        text_box_name = builder.get_object("client_name")
        text_box_build = builder.get_object("job_name")
        text_box_adress = builder.get_object("adress")
        text_box_telephone= builder.get_object("telephone")

        document = Document("./datos/plantilla.docx")

        self.insert_text_bold(document, 3, 'NOMBRE: ' , text_box_name.get_text())
        self.insert_text_bold(document, 4, 'OBRA: ' , text_box_build.get_text())
        self.insert_text_bold(document, 5, 'DIRECCIÓN: ' , text_box_adress.get_text())
        self.insert_text_bold(document, 6, 'TELÉFONO: ' , text_box_telephone.get_text())
    
        self.modify_table(document)

        document.save('./datos/new_document.docx')
        



    def button_input_mount_pressed(self, button):
        input_mount = builder.get_object("input_value")
        self.print_total(int(input_mount.get_text()),"Varios")
    def onDestroy(self, *args):
        print("Close program")
        Gtk.main_quit()

builder = Gtk.Builder()
builder.add_from_file("inform_generator.glade")
new_manager = PhotocopyManager()
handler = Handler(new_manager)
builder.connect_signals(handler)

window = builder.get_object("window1")
#window.set_icon_from_file('cat_logo.png')
window.connect("destroy",Gtk.main_quit)
window.show_all()
Gtk.main()
