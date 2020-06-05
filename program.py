import gi
gi.require_version("Gtk", "3.0")
from gi.repository import Gtk
from datetime import date
from datetime import datetime
import os.path
import platform, subprocess
import shutil
import re
from utils import formated_namber
from utils import get_text_moth
from inform_mail import send_email

from docx import Document
from photocopy import PhotocopyManager

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from num2words import num2words
from random import randint
import subprocess

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

today = date.today()


class Handler:
    manager = None 
    total = 0
    dialog = None
    inform_generated = False
    table_count = 0
    table_row_to_add = 0

    inform_id = 0

    def __init__(self, manager):
        filepath = "./datos/presupuesto_numero.txt"
        file_inform_number = open(filepath,"r")
        number_line = file_inform_number.read(4)
        self.inform_number = number_line
        self.inform_id = int(number_line)

        self.label_inform_number = builder.get_object("label_inform_number")
        self.text_box_name = builder.get_object("client_name")
        self.text_box_build = builder.get_object("job_name")
        self.text_box_adress = builder.get_object("adress")
        self.text_box_telephone= builder.get_object("telephone")
        self.description_obj = builder.get_object("in_description")
        self.price_obj = builder.get_object("in_price")
        self.count_obj = builder.get_object("in_count")
        self.input_inform_name = builder.get_object("input_inform_name")
        self.total_label = builder.get_object("label_total")

        self.payment_type = "CONTADO"

        rb_meter = builder.get_object("rb_meters")
        rb_meter2 = builder.get_object("rb_m2")
        rb_unit= builder.get_object("rb_unit")
        rb_ml= builder.get_object("rb_ml")

        rb_5050 = builder.get_object("rb_5050")
        rb_cash = builder.get_object("rb_cash")

        rb_cash.connect("toggled",self.rb_action_cash)
        rb_5050.connect("toggled",self.rb_action_5050)        

        self.label_inform_number.set_text(self.inform_number+"-")
        
        self.spin = builder.get_object("spin_id")
        
        self.inform_number = self.label_inform_number.get_text()+str(self.spin.get_value_as_int())

        self.grid = Gtk.Grid()
        
        hbox = builder.get_object("main_box")

        self.list = Gtk.ListStore(str,str,str,str,str)

        self.treeview = Gtk.TreeView.new_with_model(self.list)
        for i, column_title in enumerate(["Descripcion","Cantidad","Medida","Precio","Importe"]):
            renderer = Gtk.CellRendererText()
            column = Gtk.TreeViewColumn(column_title,renderer, text=i)
            self.treeview.append_column(column)
        
        hbox.pack_start(self.treeview,True,True,1)
        #radios buttons units
        

        rb_meter.connect("toggled",self.rb_action_meters)
        rb_meter2.connect("toggled",self.rb_action_m2)
        rb_unit.connect("toggled", self.rb_action_units)
        rb_ml.connect("toggled", self.rb_action_ml)

        self.messure = "m2" 
        
        total_label = builder.get_object("label_total")
        total_label.set_text("0")
        
        self.btn_edit = builder.get_object("btn_edit")
        self.btn_print = builder.get_object("btn_print")
        self.btn_new = builder.get_object("btn_new_inform")
        self.entry_save_as = builder.get_object("input_inform_name")
        self.btn_save_as = builder.get_object("btn_save_as")

        text_box_name = builder.get_object("client_name")
        #text_box_name.connect("insert_text",self.on_entry_insert_text)
        text_box_name.connect("changed",self.on_entry_changed)
        
        in_price = builder.get_object("in_price")
        #in_price.connect("insert_text",self.on_entry_insert_text_number)
        in_price.connect("changed",self.on_entry_number_changed)
        

    def rb_action_ml(self,button):
        self.messure = "ml"
    def rb_action_m2(self,button):
        self.messure = "m2"
    def rb_action_meters(self,button):
        self.messure = "m"
    def rb_action_units(self,button):
        self.messure = "und"
    def rb_action_5050(self,button):
        self.payment_type = "50% al empezar y el saldo al finalizar el trabajo."
    
    def rb_action_cash(self,button):
        self.payment_type = "CONTADO"

    def init(self):
        #manager.init_logger()    
        self.manager = manager 
        self.total = manager.total
        self.update_total_label()
        self.update_half_total_label()
        button1 = builder.get_object("rb_1")
        button1.connect("toggled",self.on_radio_button_marta_select)
        
        inputone = builder.get_object("input")
        self.connect('changed', self.on_changed)

        button2 = builder.get_object("rb_2")
        button3 = builder.get_object("rb_3")
        button2.connect("toggled",self.on_radio_button_oski_select)
        button3.connect("toggled",self.on_radio_button_pancha_select)

        print_black_and_white_radio = builder.get_object("black_white_radio_button")
        print_color_radio = builder.get_object("color_print_radio_button")
        print_black_and_white_radio.connect("toggled",self.radio_button_black_white_pressed)
        print_color_radio.connect("toggled",self.radio_button_color_pressed)

    def radio_button_black_white_pressed(self,button):
        self.manager.current_print_type = "black_and_white"
    def radio_button_color_pressed(self, button):
        self.manager.current_print_type = "color"

    def update_total_label(self):
        label_total = builder.get_object("label_total")
        label_total.set_text(formated_namber(self.total))
        label_total_in_the_box = builder.get_object("label_total_in_the_box")
        label_total_in_the_box.set_text(formated_namber(self.manager.initial_value_in_the_box + self.total)) 
        self.update_half_total_label()
    def print_total_to_inform_file(self):
        from_file = today_file("r") 
        line = from_file.readline()
        # make any changes to line here
        line = "Fecha: "
        line += str(today) + "                               "
        line += "TOTAL: " + formated_namber(self.total) + "\n"
        to_file = today_file("w") 
        to_file.write(line)
        shutil.copyfileobj(from_file, to_file)
        
    
    def update_half_total_label(self):
        label_half_total = builder.get_object("label_halft_total")  
        label_half_total.set_text(formated_namber(self.total/2)) 

    def print_total(self, price, data_type):
        current_log_file = open("./datos/"+str(today)+".txt","a")
        self.total = self.total + price
        self.update_total_label() 
        current_time = datetime.now().strftime("%H:%M:%S")
        formated_price =  "{:,}".format(price)
        current_log_file.write(current_time + 
                "                " + formated_price + 
                "                " + data_type + "\n" )
        current_log_file.close() 
        self.print_total_to_inform_file()
        self.manager.previous_added_price = price

    ###########################################
    ############       Buttons      ###########
    ###########################################

    def button_undo_pressed(self, button):
        print("undo") 
        opend_file = today_file("r")
        lines = opend_file.readlines()
        opend_file.close()
        new_file = today_file("w")
        for line in range(0,(len(lines)-1)):
            new_file.write(lines[line])
        new_file.close()
        self.total -= self.manager.previous_added_price
        self.print_total_to_inform_file()
        self.update_total_label()
    def button_retire_pressed(self, button):
        print("retire") 
        self.dialog = builder.get_object("retire_dialog")
        response = self.dialog.run()
    
    def on_dialog_delete_event(self, dialog, event):
        dialog.hide()
        return True

    def button_retire_accept_pressed(self, button):
        input_retire_mount = builder.get_object("dialog_mount_input")
        input_value = input_retire_mount.get_text()
        print(input_value)
        print("accept")
        
        today = date.today()
        out_log_file= open("./datos/"+str(today)+"_out"+".txt","w+")
        out_log_file.write(self.manager.current_user+": ") 
        out_log_file.write(input_value) 
        out_log_file.write("\n") 
        out_log_file.close()
        self.dialog.hide()

    def on_radio_button_select(self, widget , data=None):
        print("radio button changed")
    def on_radio_button_marta_select(self,widget):
        print("marta")
        self.manager.current_user = "Marta"

    def on_radio_button_pancha_select(self,widget):
        print("pacha")
        self.manager.current_user = "Pancha"

    def on_radio_button_oski_select(self,widget):
        print("oski")
        self.manager.current_user = "Oski"

    def button_print_pressed(self, button):
        print("printing")
        today = date.today()
        if platform.system() == 'Windows':    # Windows
            filepath = "datos/" + str(today) + ".txt" 
            relative_path = os.path.abspath(filepath) 
            os.startfile(relative_path, "print") 

    def button_cancel_clicked(self, button):
        print("cancel")
        self.dialog.hide()

    def button_show_data_pressed(self, button):
        today = date.today()
        if platform.system() == 'Windows':    # Windows
            filepath = "datos/" + str(today) + ".txt" 
            relative_path = os.path.abspath(filepath) 
            os.startfile(relative_path)
        else:
            filepath = "./datos/" + str(today) + ".txt" 
            subprocess.call(('xdg-open', filepath))

    
    def insert_text_bold(self, document, paragraph_index  ,title , value):
        document.paragraphs[paragraph_index].text = ''
        run = document.paragraphs[paragraph_index].add_run(title)
        run.font.bold = True
        run.font.name = "Arial"
        run = document.paragraphs[paragraph_index].add_run(value)
        run.font.bold = True
        run.font.name = "Arial"
    
    def table_clean(table):
        for row in table.rows:
            row.cells[0].text = ""
    def table_get_row_count(self, table):
        row_count = 0
        for row in table.rows:
            row_count = row_count + 1
        return row_count

    def set_cell_border(self, cell: _Cell, **kwargs):
        """
        Set cell`s border
        Usage:

        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('right', 'top', 'left', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))


    def modifyBorder(self,table):
        tbl = table._tbl # get xml element in table
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
        
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
        
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'auto')

            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'single')

            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)
    
    def modify_table(self, document):
        print("table")
        i = 1 
        table = document.tables[0]
        
        for elem in self.list:
            table.add_row() 
            (des , count , unit ,price , total) = elem   
            run = table.cell(i,0).paragraphs[0].add_run(des)
            run.font.size = Pt(10)
            run.font.name = "Arial"  

            count_text = str(count)+unit+"."          
            run = table.cell(i,1).paragraphs[0].add_run(count_text)
            run.font.size = Pt(10)
            run.font.name = "Arial" 
            table.cell(i,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            price_text = str(price)
            run = table.cell(i,2).paragraphs[0].add_run(price_text)
            run.font.size = Pt(10)
            run.font.name = "Arial" 
            table.cell(i,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run = table.cell(i,3).paragraphs[0].add_run(str(total))
            run.font.size = Pt(10)
            run.font.name = "Arial"            
            table.cell(i,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            self.set_cell_border(
                    table.cell(i,3),
                       right={"sz": 12, "val": "single"},
                       left={"sz": 12, "val": "single"}       
            )
            self.set_cell_border(
                    table.cell(i,0),
                       left={"sz": 12, "val": "single"},       
            )


            i = i + 1
        
        table.add_row()
        table.add_row()

        row_count = self.table_get_row_count(table)

        run = table.cell(row_count-2,0).paragraphs[0].add_run("TOTAL: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"

        run = table.cell(row_count-2,3).paragraphs[0].add_run(formated_namber(self.total))
        run.font.size = Pt(10)
        run.font.name = "Arial"
        table.cell(row_count-2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT 

        text_total_line = "Son Gs.: " + num2words(self.total,lang='es') + "."
        run = table.cell(row_count-1,0).paragraphs[0].add_run(text_total_line)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(10)
       
        self.modifyBorder(table)

        total_cell = table.cell(row_count-2,0)
        
        text_total_cell = table.cell(row_count-1,0)
        text_total_cell_nil = table.cell(row_count-1,3)
        total_cell_end= table.cell(row_count-2,3)

        self.set_cell_border(
            total_cell,
            top={"sz": 12, "val": "single"},           
            bottom={"sz": 12, "val": "single"},   
            left={"sz": 12, "val": "single"}, 
            right={"sz": 12, "val": "single"},       
        )

        self.set_cell_border(
            text_total_cell,
            top={"sz": 12, "val": "single"},           
            bottom={"sz": 12, "val": "single"},   
            left={"sz": 12, "val": "single"}, 
            right={"sz": 12, "val": "single"},       
        )

        self.set_cell_border(
            text_total_cell_nil,
            top={"sz": 12, "val": "nil"},           
            bottom={"sz": 12, "val": "nil"},   
            left={"sz": 12, "val": "nil"}, 
            right={"sz": 12, "val": "nil"},       
        )

        
        self.set_cell_border(
            total_cell_end,
            top={"sz": 12, "val": "single"},           
            bottom={"sz": 12, "val": "single"},   
            left={"sz": 12, "val": "single"}, 
            right={"sz": 12, "val": "single"},       
        )
  
        cell1 = table.cell(row_count-1,0)
        cell2 = table.cell(row_count-1,1)
        cell3 = table.cell(row_count-1,2)
       
        cell1.merge(cell2)
        cell1.merge(cell3) 

        cell1 = table.cell(row_count-2,0)
        cell2 = table.cell(row_count-2,1)
        cell3 = table.cell(row_count-2,2)

        cell1.merge(cell2)
        cell1.merge(cell3) 

    def on_button_generate_pressed(self, button):
        print("generate")
        self.inform_generated = True
     

        day = today.strftime("%d")
        year = today.strftime("%Y")
        moth_number = int(today.strftime("%m"))
        inform_date = "Encarnacion " + day + " de " + get_text_moth(moth_number) + " de " + year          
        inform_date = inform_date.upper()

        document = Document("./datos/plantilla.docx")

        self.insert_text_bold(document, 0, 'PRESUPUESTO' , "                                           "+ "N°: " + self.inform_number)
        
        self.insert_text_bold(document, 2, inform_date , '')

        self.insert_text_bold(document, 3, 'NOMBRE: ' , self.text_box_name.get_text())
        self.insert_text_bold(document, 4, 'OBRA: ' , self.text_box_build.get_text())
        self.insert_text_bold(document, 5, 'DIRECCIÓN: ' , self.text_box_adress.get_text())
        self.insert_text_bold(document, 6, 'TELÉFONO: ' , self.text_box_telephone.get_text())
    

        paragraph_index = 14
        document.paragraphs[paragraph_index].text = ''
        run = document.paragraphs[paragraph_index].add_run("FORMA DE PAGO: ")
        run.font.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run = document.paragraphs[paragraph_index].add_run(self.payment_type)
        run.font.name = "Arial"
        run.font.size = Pt(10)


        self.modify_table(document)
       
        self.btn_edit.set_visible(True)
        self.btn_print.set_visible(True)

        self.btn_new.set_visible(True)
        self.entry_save_as.set_visible(True)
        self.btn_save_as.set_visible(True)

        document.save('./datos/presupuesto_generado.docx')
        self.document = document
        
    def btn_delete(self, button):
        print("delete")
        selected = self.treeview.get_selection()
        (model, paths) = selected.get_selected_rows()
        for path in paths:
           iter = model.get_iter(path)
           text_value = self.list[iter][4]
           value = text_value.replace('.','')
           print(int(value))
           self.total = self.total - int(value) 
           total_label = builder.get_object("label_total")
           total_label.set_text(formated_namber(self.total))
           model.remove(iter)


    def btn_add(self, button):
        print("add")
        self.table_count = self.table_count + 1 
        if(self.table_count > 5):
            self.table_row_to_add = self.table_row_to_add + 1
        
        count = float(self.count_obj.get_text())
        price = self.price_obj.get_text().replace('.','')
        price = int(price)
        import_value = int(price * count)
        new_element = (self.description_obj.get_text(), self.count_obj.get_text() ,self.messure, formated_namber(price) , formated_namber(import_value)) 
        self.list.append(list(new_element))
        self.total = self.total + import_value         
       
        self.total_label.set_text(formated_namber(self.total))
        self.price_obj.set_text("")
        self.count_obj.set_text("")
        self.description_obj.set_text("")

    def button_input_mount_pressed(self, button):
        input_mount = builder.get_object("input_value")
        self.print_total(int(input_mount.get_text()),"Varios")
    def onDestroy(self, *args):
        print("Close program")
        Gtk.main_quit()

    def btn_print_clicked(self, button):
        print("printing")
        if platform.system() == 'Windows':    # Windows
            path = "./datos/presupuesto_generado.docx"
            #relative_path = os.path.abspath(path) 
            #A
            # os.startfile(relative_path, "print") 
            subprocess.call("printing.bat")
        else:
            print("print in freebsd")

    def btn_inform_show(self, button):
        if(self.inform_generated == False):
            print("No file generated") 
            return
        if platform.system() == 'Windows':    # Windows
            filepath = "datos/presupuesto_generado.docx" 
            relative_path = os.path.abspath(filepath) 
            os.startfile(relative_path)
        else:
            filepath = "datos/presupuesto_generado.docx" 
            subprocess.call(('xdg-open', filepath))

    def btn_save_as_clicked(self, button):
        print("save as")
        entry = builder.get_object("input_inform_name")
        self.document.save("./datos/"+entry.get_text())
        
    def on_entry_changed(self,entry):
        input_save_name = builder.get_object("input_inform_name")
        self.inform_number = self.label_inform_number.get_text()+str(self.spin.get_value_as_int())
        input_save_name.set_text(self.inform_number+entry.get_text()+".docx")

    def on_entry_insert_text(self,entry,text,length,position):
        print("insert")
    def btn_send_mail(self, button):
        print("email")
        input_mail = builder.get_object("input_send_mail") 
        
        send_email(input_mail.get_text(),
           'Presupuesto',
           'Presupuesto', 
           'datos/presupuesto_generado.docx')
    def btn_new_clicked(self, button):
        print("new")
        self.inform_id += 1
        print(self.inform_id)
        
        self.label_inform_number = builder.get_object("label_inform_number")
        self.label_inform_number.set_text(str(self.inform_id)+"-")
        open_file = open("./datos/presupuesto_numero.txt","w+")
        open_file.write(str(self.inform_id))  
        self.list.clear()
        self.price_obj.set_text("")
        self.count_obj.set_text("")
        self.description_obj.set_text("")
        self.text_box_name.set_text("") 
        self.text_box_build.set_text("")
        self.text_box_adress.set_text("")
        self.text_box_telephone.set_text("")
        self.total = 0
        self.total_label.set_text(formated_namber(self.total))
        self.table_row_to_add = 0  
        self.table_count = 0
        self.input_inform_name.set_text('')

    def on_entry_number_changed(self,entry):
        formated = "{:,}".format(entry.get_text())
        formated = formated.replace(',','.')
        entry.set_text(formated)
    def on_entry_insert_text_number(self,entry,text,length,position):
         
        print("insert number")

builder = Gtk.Builder()
builder.add_from_file("inform_generator.glade")
new_manager = PhotocopyManager()
handler = Handler(new_manager)
builder.connect_signals(handler)
window = builder.get_object("window1")
#window.set_icon_from_file('cat_logo.png')
window.connect("destroy",Gtk.main_quit)
window.show_all()
handler.btn_edit.set_visible(False)
handler.btn_print.set_visible(False)
handler.btn_new.set_visible(False)
handler.entry_save_as.set_visible(False)
handler.btn_save_as.set_visible(False)
Gtk.main()
